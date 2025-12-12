import os
import re
import eel
import pyautogui
import screen_brightness_control as sbc
from AppOpener import open as appopen
import pywhatkit
import google.generativeai as genai
import psutil
import requests
import hashlib
from bs4 import BeautifulSoup
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from engine.config import ASSISTANT_NAME, GEMINI_API_KEY

# Gemini Setup - USING STABLE 1.5 FLASH MODEL
genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-2.0-flash') 
chat_session = model.start_chat(history=[
    {"role": "user", "parts": ["You are Vaytron. Answer short and concisely."]},
    {"role": "model", "parts": ["Okay."]}
])

# --- HELPER: Safe Speak Function ---
def speak(text):
    from engine.command import speak as sp
    sp(text)

# --- FEATURES ---

def openCommand(query):
    query = query.lower().replace("open", "").replace("launch", "").strip()
    if isinstance(ASSISTANT_NAME, list):
        for name in ASSISTANT_NAME: query = query.replace(name.lower(), "")
    else: query = query.replace(ASSISTANT_NAME.lower(), "")
    
    if query != "":
        speak("Opening " + query)
        try: appopen(query, match_closest=True, throw_error=True)
        except: os.system('start ' + query)

def setVolume(command):
    from engine.command import speak
    if "up" in command or "increase" in command:
        pyautogui.press("volumeup")
        speak("Volume increased")
    elif "down" in command or "decrease" in command:
        pyautogui.press("volumedown")
        speak("Volume decreased")
    elif "mute" in command:
        pyautogui.press("volumemute")
        speak("System muted")

def setBrightness(command):
    from engine.command import speak
    current = sbc.get_brightness()
    if not current: current = [50]
    
    if "up" in command or "increase" in command:
        sbc.set_brightness(min(current[0] + 10, 100))
        speak("Brightness increased")
    elif "down" in command or "decrease" in command:
        sbc.set_brightness(max(current[0] - 10, 0))
        speak("Brightness decreased")

def sendWhatsApp(query):
    from engine.command import speak
    contacts = {"mom": "+910000000000", "dad": "+910000000000"} 
    for name, number in contacts.items():
        if name in query:
            msg = query.replace("send message to", "").replace(name, "").replace("saying", "").strip()
            speak(f"Sending message to {name}")
            pywhatkit.sendwhatmsg_instantly(number, msg, wait_time=10)
            return
    speak("I couldn't find that contact.")

def system_stats():
    cpu = psutil.cpu_percent()
    battery = psutil.sensors_battery()
    msg = f"CPU is at {cpu} percent."
    if battery: msg += f" Battery is at {battery.percent} percent."
    speak(msg)

def check_password_safety(password):
    sha1 = hashlib.sha1(password.encode('utf-8')).hexdigest().upper()
    try:
        res = requests.get('https://api.pwnedpasswords.com/range/' + sha1[:5], timeout=5)
        if sha1[5:] in res.text: speak("Alert! Password found in data breaches.")
        else: speak("Password seems safe.")
    except: speak("Could not check password database.")

def research_topic(topic):
    speak(f"Researching {topic}...")
    try:
        res = requests.get(f"https://www.google.com/search?q={topic}", headers={'User-Agent': 'Mozilla/5.0'})
        soup = BeautifulSoup(res.text, 'html.parser')
        results = [r.get_text() for r in soup.find_all('div', class_='BNeawe s3v9rd AP7Wnd')[:5]]
        
        doc = Document()
        doc.add_heading(topic, 0)
        for line in results: doc.add_paragraph(line)
        filename = f"{topic.replace(' ', '_')}.docx"
        doc.save(filename)
        speak("Report saved.")
        os.system(f"start {filename}")
    except: speak("Research failed.")

def open_social_media(platform):
    speak(f"Opening {platform}...")
    options = webdriver.ChromeOptions()
    options.add_experimental_option("detach", True)
    try:
        driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
        driver.get(f"https://{platform}.com")
    except: speak("Could not open browser.")

# --- CHAT LOGIC (ONLINE ONLY) ---
def chatWithBot(query):
    eel.DisplayMessage("Thinking...")
    
    try:
        response = chat_session.send_message(query)
        speak(response.text.replace("*", ""))
    except Exception as e:
        err_msg = str(e)
        if "429" in err_msg:
            speak("I am tired. My API quota is full. Please wait a minute.")
        else:
            print(f"Error: {e}")
            speak("I am unable to connect to the internet right now.")